"""
Week 5 Report Generator – Shruti Malik
Generates a Word document (.docx) for the analytics project draft
"""

import os
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import datetime

BASE_DIR = r"C:\Users\Owner\Documents\Analytics-Project-Draft"

def set_font(run, name="Times New Roman", size=11, bold=False):
    run.font.name = name
    run.font.size = Pt(size)
    run.font.bold = bold

def add_heading(doc, text, level=1):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    run.font.name = "Times New Roman"
    run.font.size = Pt(12 if level == 1 else 11)
    run.font.bold = True
    return p

def add_body(doc, text, indent=False):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.first_line_indent = Inches(0.5)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    run = p.add_run(text)
    set_font(run)
    return p

def add_figure(doc, path, caption):
    if os.path.exists(path):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.add_picture(path, width=Inches(5.5))
        cap = doc.add_paragraph(caption)
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cap.paragraph_format.space_after = Pt(10)
        r = cap.runs[0] if cap.runs else cap.add_run(caption)
        r.font.italic = True
        r.font.size = Pt(9)
        r.font.name = "Times New Roman"

doc = Document()

# ── Page margins ──────────────────────────────────────────────────────────────
for section in doc.sections:
    section.top_margin    = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin   = Inches(1)
    section.right_margin  = Inches(1)

# ══════════════════════════════════════════════════════════════════════════════
# TITLE PAGE
# ══════════════════════════════════════════════════════════════════════════════
for _ in range(4):
    doc.add_paragraph()

title_p = doc.add_paragraph()
title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = title_p.add_run("Predicting Prescription Non-Adherence Through Hospital Readmission Patterns:\nAn Analytic Project Report Draft")
r.font.name = "Times New Roman"; r.font.size = Pt(14); r.font.bold = True

for _ in range(2):
    doc.add_paragraph()

for line in ["Shruti Malik", "ITEC 5040: Predictive Analytics",
             "Capella University", "Dr. C", "May 2026"]:
    p = doc.add_paragraph(line)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.runs[0]
    r.font.name = "Times New Roman"; r.font.size = Pt(11)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# ABSTRACT
# ══════════════════════════════════════════════════════════════════════════════
add_heading(doc, "Abstract")
add_body(doc, (
    "This report summarizes the analytical work completed to date on an individual course project "
    "examining prescription non-adherence among diabetic patients through the lens of hospital "
    "readmission patterns. Using the publicly available Diabetes 130-US Hospitals dataset (1999–2008) "
    "from the UCI Machine Learning Repository, this paper defines the core business problem, evaluates "
    "data alignment, addresses ethical and legal considerations, presents a comprehensive data audit "
    "with visualizations, and documents all data cleansing procedures with accompanying code. "
    "The cleaned dataset, comprising 69,987 unique patient encounters across 46 features, provides a "
    "strong analytical foundation for the classification modeling work planned in subsequent project phases."
))
doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 1 – BUSINESS QUESTION
# ══════════════════════════════════════════════════════════════════════════════
add_heading(doc, "Business Question, Problem, and Data Source")
add_body(doc, (
    "Prescription abandonment — the act of dropping off a prescription at a pharmacy and never returning "
    "to pick it up — is a pervasive and underexamined failure point in the U.S. medication adherence "
    "pipeline. Studies estimate that approximately 20–30% of new prescriptions are never filled, and "
    "nearly 50% of medications for chronic conditions are not taken as prescribed (Kleinsinger, 2018). "
    "For patients managing diabetes, a condition requiring sustained, multi-drug regimens, non-adherence "
    "carries severe downstream consequences: uncontrolled blood glucose, accelerated comorbidity "
    "progression, preventable emergency department visits, and costly hospital readmissions."
))
add_body(doc, (
    "Drawing on several years of experience as a pharmacy technician, I observed this phenomenon "
    "firsthand. Prescriptions returned to stock after patients failed to retrieve them were not random — "
    "they clustered around certain patient profiles: older adults managing multiple chronic conditions, "
    "recently discharged patients overwhelmed by new medication regimens, and those whose medications "
    "had changed during a hospitalization. The business question driving this project is: Can patient-level "
    "clinical and demographic data collected at the time of hospital discharge be used to predict the "
    "likelihood of medication non-adherence and, by proxy, early hospital readmission in diabetic patients? "
    "An accurate predictive model would allow pharmacies and care coordinators to flag high-risk patients "
    "for proactive outreach before prescription abandonment occurs."
))

add_heading(doc, "Data Source", level=2)
add_body(doc, (
    "The primary dataset for this project is the Diabetes 130-US Hospitals for Years 1999–2008 dataset, "
    "hosted by the UCI Machine Learning Repository (Strack et al., 2014). The dataset was originally "
    "compiled from the Health Facts database maintained by Cerner Corporation and represents clinical "
    "records from 130 U.S. hospitals. It contains 101,766 inpatient encounters and 50 features spanning "
    "patient demographics (age, race, gender), administrative variables (admission type, discharge "
    "disposition, length of stay), laboratory results (HbA1c levels, serum glucose measurements), "
    "ICD-9 primary and secondary diagnoses, counts of procedures and medications, changes in diabetic "
    "drug regimens, and the key target variable: readmission status, coded as not readmitted (NO), "
    "readmitted after 30 days (>30), or readmitted within 30 days (<30)."
))

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 2 – DATA ALIGNMENT
# ══════════════════════════════════════════════════════════════════════════════
add_heading(doc, "Data Alignment With the Business Problem")
add_body(doc, (
    "At first consideration, a hospital readmission dataset and a retail pharmacy abandonment problem may "
    "appear conceptually misaligned. Hospital readmission is an outcome observed weeks after a patient "
    "leaves a facility, while prescription abandonment occurs within hours of discharge. However, the "
    "alignment is both intentional and well-supported by the literature. Lam and Marsden (2014) "
    "established that medication non-adherence is one of the strongest independent predictors of "
    "30-day readmission. In this framework, early readmission functions as a lagged, observable signal "
    "of the same underlying behavior: a patient who did not follow through with their medication plan."
))
add_body(doc, (
    "The features available in the dataset map directly onto the patient-level risk factors I observed "
    "empirically in a pharmacy setting. The number of medications variable captures polypharmacy burden — "
    "a documented driver of adherence fatigue. The change in diabetic medications indicator identifies "
    "patients whose regimens were altered during the admission, a known moment of heightened "
    "non-adherence risk. The number of prior inpatient visits reflects the chronic hospitalization "
    "patterns of patients with low outpatient engagement. Age and race encode socioeconomic and "
    "health literacy dimensions that pharmacy staff often navigate informally. Primary diagnosis codes "
    "allow stratification by comorbidity complexity. Together, these variables make the dataset a "
    "reasonable and well-motivated proxy data source for predicting prescription non-adherence."
))
add_body(doc, (
    "One area of partial misalignment merits acknowledgment: the dataset does not contain any explicit "
    "pharmacy-side variables such as prescription fill rates, time-to-pick-up, or copay information. "
    "These variables would ideally be present in a purpose-built adherence dataset. To address this gap, "
    "I will treat the binary readmission outcome (readmitted within 30 days vs. not) as the operational "
    "proxy for non-adherence and will limit my interpretation of model outputs to identifying high-risk "
    "patient profiles rather than making direct causal claims about pharmacy behavior. This framing is "
    "consistent with how similar datasets have been applied in the academic adherence literature "
    "(Strack et al., 2014)."
))

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 3 – ETHICAL, LEGAL, GLOBAL, AND CULTURAL CONSIDERATIONS
# ══════════════════════════════════════════════════════════════════════════════
add_heading(doc, "Ethical, Legal, Global, and Cultural Considerations")
add_body(doc, (
    "Health data carries a unique constellation of ethical and legal obligations that must be explicitly "
    "addressed before any analytical work proceeds. The considerations most relevant to this project span "
    "four domains: privacy and HIPAA compliance, algorithmic fairness, global applicability, and cultural "
    "sensitivity in model interpretation."
))
add_body(doc, (
    "Privacy and Legal Compliance. The Diabetes 130-US Hospitals dataset is fully de-identified and "
    "publicly available through the UCI Machine Learning Repository. No patient names, social security "
    "numbers, dates of birth, geographic identifiers smaller than a state, or other Protected Health "
    "Information (PHI) as defined by the Health Insurance Portability and Accountability Act of 1996 "
    "(HIPAA) are present in the data. Because this project uses only de-identified, publicly released "
    "data for academic purposes, it does not require Institutional Review Board (IRB) oversight or data "
    "use agreements. There is no proprietary organizational data involved at any stage."
))
add_body(doc, (
    "Algorithmic Fairness and Racial Bias. The dataset includes self-reported race as a variable, with "
    "categories including Caucasian, African American, Hispanic, Asian, and Other. Incorporating race as "
    "a predictive feature introduces a documented risk of encoding structural health disparities into the "
    "model rather than illuminating them. Research has shown that algorithmic health risk scores trained "
    "on biased historical data can perpetuate unequal care allocation (Obermeyer et al., 2019). For this "
    "project, race will be treated as a fairness audit variable rather than a predictive input. Model "
    "performance will be evaluated across racial subgroups to ensure that the classifier does not "
    "systematically misclassify patients from any demographic group."
))
add_body(doc, (
    "Global and Cultural Applicability. The dataset was collected exclusively from U.S. hospitals between "
    "1999 and 2008, a period during which diabetes management protocols, insurance structures, and "
    "medication formularies differed substantially from contemporary standards. Medications like GLP-1 "
    "agonists, which now constitute a cornerstone of diabetes management, were not widely used during the "
    "data collection period. Findings from this model should not be generalized to international "
    "healthcare contexts or applied as policy guidance without substantial re-validation on more recent, "
    "culturally relevant data."
))
add_body(doc, (
    "Cultural Sensitivity. Diabetic management is deeply influenced by cultural attitudes toward food, "
    "medication, and healthcare engagement. Disparities in health literacy and access to pharmacy "
    "services are particularly pronounced in Hispanic and African American communities, which are "
    "disproportionately represented among diabetic patients. Any model outputs or patient outreach "
    "strategies developed from this work must be designed with these cultural dimensions in mind, "
    "ensuring that interventions are linguistically accessible and structurally equitable."
))

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 4 – DATA AUDIT AND PROFILING
# ══════════════════════════════════════════════════════════════════════════════
add_heading(doc, "Data Audit and Profiling")
add_body(doc, (
    "The raw dataset was loaded into Python using the pandas library, with the '?' character specified "
    "as a null value upon import. The dataset contains 101,766 rows and 50 columns, encompassing a mix "
    "of integer and object (string/categorical) data types. An initial profiling pass was conducted to "
    "assess missingness, data type integrity, distributional characteristics, and target variable balance."
))

add_heading(doc, "Missing Value Analysis", level=2)
add_body(doc, (
    "Nine of the 50 columns contained missing values. The severity of missingness varied dramatically "
    "across columns. The weight column was missing for 98,569 encounters (96.86%), effectively rendering "
    "it analytically unusable. The max_glu_serum column (fasting serum glucose) was missing for 94.75% "
    "of records, and A1Cresult was absent for 83.28% — a surprising gap given that HbA1c measurement is "
    "central to the original research paper accompanying this dataset (Strack et al., 2014). The "
    "medical_specialty column was missing for 49.08% of encounters. These four columns exceed the "
    "threshold at which imputation is reliable and must be handled through either exclusion or careful "
    "surrogate modeling. Figure 1 illustrates the missing value profile across all affected columns."
))
add_figure(doc,
    os.path.join(BASE_DIR, "fig1_missing_values.png"),
    "Figure 1. Missing value percentages across features with at least one missing record.")

add_heading(doc, "Target Variable Distribution", level=2)
add_body(doc, (
    "The target variable, readmitted, is categorical with three levels: NO (not readmitted), >30 "
    "(readmitted after 30 days), and <30 (readmitted within 30 days). The distribution is heavily "
    "imbalanced: 53.9% of encounters resulted in no readmission, 34.9% were readmitted after 30 days, "
    "and only 11.2% were readmitted within 30 days. This class imbalance is expected given the clinical "
    "reality but presents a modeling challenge that will be addressed in subsequent project phases "
    "through techniques such as stratified sampling, SMOTE oversampling, or class-weight adjustment. "
    "Figure 2 displays this distribution."
))
add_figure(doc,
    os.path.join(BASE_DIR, "fig2_readmission_dist.png"),
    "Figure 2. Distribution of the readmission target variable across 101,766 encounters.")

add_heading(doc, "Age Distribution", level=2)
add_body(doc, (
    "Patient ages are encoded as decade-width brackets. As shown in Figure 3, the patient population is "
    "heavily skewed toward older adults, with the 70–80 age bracket representing the largest single "
    "group. This is consistent with the epidemiological profile of inpatient diabetic populations and "
    "aligns with the pharmacy observation that elderly patients managing multiple chronic conditions "
    "represent the highest-risk segment for medication non-adherence."
))
add_figure(doc,
    os.path.join(BASE_DIR, "fig3_age_distribution.png"),
    "Figure 3. Distribution of patient encounter counts across age brackets.")

add_heading(doc, "Medication Burden", level=2)
add_body(doc, (
    "The num_medications variable ranges from 1 to 81, with a median of 15 medications per encounter "
    "(Figure 5). High medication counts are a known predictor of adherence breakdown; patients managing "
    "15 or more concurrent medications face compounded cognitive load, cost burdens, and scheduling "
    "complexity. The distribution is right-skewed with a long tail, indicating that while most patients "
    "manage 10–20 medications, a subset of high-complexity patients manage 40 or more."
))
add_figure(doc,
    os.path.join(BASE_DIR, "fig5_num_medications.png"),
    "Figure 5. Distribution of the number of medications per encounter (median = 15).")

add_heading(doc, "Readmission Rate by Age Group", level=2)
add_body(doc, (
    "Figure 6 plots the 30-day readmission rate by age group. A U-shaped pattern emerges: patients in "
    "the youngest brackets (0–30) and the oldest brackets (80–100) show the highest early readmission "
    "rates, while middle-aged adults (40–70) exhibit lower rates. This non-linear relationship with age "
    "suggests that any model relying on a linear encoding of age may underfit; polynomial or categorical "
    "treatment of the age feature may be warranted."
))
add_figure(doc,
    os.path.join(BASE_DIR, "fig6_readmit_by_age.png"),
    "Figure 6. Thirty-day readmission rate (%) by patient age bracket.")

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 5 – DATA CLEANSING
# ══════════════════════════════════════════════════════════════════════════════
add_heading(doc, "Addressing Missing Values and Data Issues")
add_body(doc, (
    "Data cleansing followed a structured, reproducible protocol implemented in Python. Each decision "
    "is grounded in analytical reasoning and documented below with the corresponding code."
))

steps = [
    ("Step A: Duplicate Encounter Removal",
     "Encounter IDs are designed to be unique. A check for duplicate encounter_id values confirmed that "
     "zero duplicate rows were present in the raw dataset, meaning all 101,766 records represent "
     "distinct hospital visits."),
    ("Step B: Invalid Gender Removal",
     "Three records contained the value 'Unknown/Invalid' in the gender column. As this represents a "
     "negligible fraction of data (<0.01%) and the value carries no interpretable analytical meaning, "
     "these rows were dropped from the dataset."),
    ("Step C: High-Missingness Column Removal",
     "Columns with more than 40% missing values — weight (96.86%), max_glu_serum (94.75%), "
     "A1Cresult (83.28%), and medical_specialty (49.08%) — were removed from the analytical dataset. "
     "The threshold of 40% is a widely cited heuristic beyond which imputation introduces more bias "
     "than the retained information offsets (van Buuren, 2018)."),
    ("Step D: Payer Code Imputation",
     "The payer_code column retained a 39.56% missing rate after the column-level drop. Because payer "
     "information may correlate with insurance coverage and socioeconomic status, it was retained and "
     "missing values were imputed using the column mode (MC = Medicare), which represented the most "
     "common payer type in the dataset."),
    ("Step E: Binary Target Construction",
     "The three-class readmitted variable was converted to a binary target: readmit_30 = 1 if the "
     "patient was readmitted within 30 days (<30), and 0 otherwise. This framing aligns with the "
     "clinical definition used by the Centers for Medicare and Medicaid Services (CMS) for hospital "
     "quality scoring and is consistent with the primary research question."),
    ("Step F: Age Bracket to Numeric Midpoint",
     "The age column contains bracket strings (e.g., '[50-60)'). These were mapped to their numeric "
     "midpoints (e.g., 55) to produce an ordinal numeric feature suitable for model training without "
     "requiring one-hot encoding of ten categories."),
    ("Step G: Removal of Deceased and Hospice Encounters",
     "Encounters with discharge_disposition_id values corresponding to deceased patients (11) or "
     "transfers to hospice facilities (13, 14, 19, 20, 21) were removed. These patients, by definition, "
     "cannot be readmitted and should not be scored by a readmission model. This step removed 2,423 "
     "records."),
    ("Step H: First-Encounter-Per-Patient Deduplication",
     "The dataset contains multiple encounters for many patients. Including multiple encounters from "
     "the same individual creates data leakage — a model trained on encounter n could implicitly learn "
     "information about encounter n+1 for the same patient. Following the methodology of Strack et al. "
     "(2014), only the first (earliest encounter_id) encounter per patient_nbr was retained. This step "
     "removed 29,353 records, yielding a final analytical dataset of 69,987 unique patient encounters "
     "across 46 features."),
]

for title, body in steps:
    add_heading(doc, title, level=2)
    add_body(doc, body)

add_body(doc, (
    "The complete Python implementation of all data cleansing steps is included in Appendix A. "
    "The cleaned dataset was saved as diabetic_data_clean.csv and will serve as the input for all "
    "subsequent modeling work in this project."
))

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 6 – SYNTHESIS AND KEY THEMES
# ══════════════════════════════════════════════════════════════════════════════
add_heading(doc, "Synthesis of Key Themes and Findings")
add_body(doc, (
    "Several converging themes emerge from the data profiling and cleansing work completed to this point. "
    "First, the dataset's heavy skew toward older adult patients reinforces the pharmacy-side observation "
    "that the highest-risk adherence population is concentrated in the 60–80 age range, where polypharmacy "
    "burdens are greatest and outpatient support systems are often weakest. The median of 15 medications "
    "per encounter is striking; no individual can reasonably be expected to self-manage that regimen "
    "without structural support."
))
add_body(doc, (
    "Second, the pervasive missingness of clinically important variables — particularly HbA1c and body "
    "weight — reveals a systemic data quality issue in the underlying health records system. The fact "
    "that over 83% of encounters lack an HbA1c result is itself a clinical finding: it suggests that "
    "glycemic testing was inconsistently applied across the 130 participating hospitals, potentially "
    "reflecting resource disparities or documentation culture differences across institutions. This "
    "finding is consistent with the central argument of Strack et al. (2014), who found that HbA1c "
    "measurement, when it did occur, was associated with meaningfully lower readmission rates."
))
add_body(doc, (
    "Third, the class imbalance in the target variable (11.2% 30-day readmissions) reflects the real-world "
    "signal-to-noise ratio of pharmacy-relevant events. Prescription abandonment and early readmission "
    "are low-frequency but high-impact events. This creates a modeling environment where accuracy is "
    "a misleading performance metric and where precision-recall tradeoffs, area under the ROC curve, "
    "and F1 scores will be more meaningful. These considerations will guide model selection and "
    "evaluation in the weeks ahead."
))
add_body(doc, (
    "Finally, the convergence of multiple evidence streams — clinical literature on adherence, "
    "pharmacist-level observational experience, and the distributional characteristics of the dataset "
    "itself — builds a robust case that the chosen data source, while not a perfect match for the "
    "business problem, is analytically appropriate and practically useful. The proxy variable "
    "approach (readmission as a signal of non-adherence) is a well-established technique in "
    "health analytics when ground-truth adherence data is unavailable."
))

# ══════════════════════════════════════════════════════════════════════════════
# CONCLUSION
# ══════════════════════════════════════════════════════════════════════════════
add_heading(doc, "Conclusion")
add_body(doc, (
    "This report summarizes the foundational work completed on the prescription non-adherence prediction "
    "project. The business problem has been clearly defined, the data source has been documented and "
    "evaluated for alignment, ethical and legal risks have been catalogued and mitigated, and a "
    "thorough data audit and cleansing protocol has been executed and documented. The resulting "
    "analytical dataset of 69,987 unique patient encounters across 46 features represents a solid "
    "and defensible foundation for the classification modeling work that will follow in subsequent "
    "project phases. The model, once complete, will produce patient-level risk scores that could be "
    "integrated into pharmacy workflows to trigger proactive outreach — turning the data-driven insight "
    "of this project into a tangible tool for improving medication adherence outcomes."
))

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# REFERENCES
# ══════════════════════════════════════════════════════════════════════════════
add_heading(doc, "References")
refs = [
    "Kleinsinger, F. (2018). The unmet challenge of medication nonadherence. The Permanente Journal, 22, 18–033. https://doi.org/10.7812/TPP/18-033",
    "Lam, W. Y., & Marsden, P. (2014). Medication adherence measures: An overview. BioMed Research International, 2015, 217047. https://doi.org/10.1155/2015/217047",
    "Obermeyer, Z., Powers, B., Vogeli, C., & Mullainathan, S. (2019). Dissecting racial bias in an algorithm used to manage the health of populations. Science, 366(6464), 447–453. https://doi.org/10.1126/science.aax2342",
    "Strack, B., DeShazo, J. P., Gennings, C., Olmo, J. L., Ventura, S., Cios, K. J., & Clore, J. N. (2014). Impact of HbA1c measurement on hospital readmission rates: Analysis of 70,000 clinical database patient records. BioMed Research International, 2014, 781670. https://doi.org/10.1155/2014/781670",
    "UCI Machine Learning Repository. (2014). Diabetes 130-US hospitals for years 1999–2008. University of California, Irvine. https://archive.ics.uci.edu/ml/datasets/Diabetes+130-US+hospitals+for+years+1999-2008",
    "van Buuren, S. (2018). Flexible imputation of missing data (2nd ed.). CRC Press. https://stefvanbuuren.name/fimd/",
]
for ref in refs:
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Inches(-0.5)
    p.paragraph_format.left_indent = Inches(0.5)
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run(ref)
    r.font.name = "Times New Roman"
    r.font.size = Pt(11)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# APPENDIX A – PYTHON CODE
# ══════════════════════════════════════════════════════════════════════════════
add_heading(doc, "Appendix A: Python Data Profiling and Cleansing Code")

code_lines = open(os.path.join(BASE_DIR, "data_profiling.py"), encoding="utf-8").readlines()
code_text  = "".join(code_lines)

p = doc.add_paragraph()
r = p.add_run(code_text)
r.font.name  = "Courier New"
r.font.size  = Pt(8)

# ── Save ──────────────────────────────────────────────────────────────────────
out_path = os.path.join(BASE_DIR, "Week5_Analytics_Report_Shruti_Malik.docx")
doc.save(out_path)
print(f"Saved: {out_path}")
